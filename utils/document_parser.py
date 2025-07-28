"""
Document parsing utilities for media plan validation
"""

import re
from typing import List, Dict, Any, Tuple, Optional
from docx import Document
from config.validation_config import (
    CURRENCY_PATTERNS, PLATFORM_KEYWORDS, OBJECTIVE_KEYWORDS,
    STRATEGY_EXPLAINER_PATTERNS, CHANNEL_NORMALIZATIONS
)


class DocumentSectionDetector:
    """Utility class for detecting sections in Word documents"""
    
    @staticmethod
    def extract_campaign_dates(doc: Document) -> Tuple[Optional[str], Optional[str]]:
        """Extract campaign start and end dates from Word document"""
        start_date = None
        end_date = None
        
        # Look through paragraphs for date patterns
        date_patterns = [
            r'campaign\s+start.*?(\d{4}-\d{2}-\d{2})',
            r'start\s+date.*?(\d{4}-\d{2}-\d{2})',
            r'campaign\s+end.*?(\d{4}-\d{2}-\d{2})',
            r'end\s+date.*?(\d{4}-\d{2}-\d{2})',
            r'duration.*?(\d{4}-\d{2}-\d{2}).*?to.*?(\d{4}-\d{2}-\d{2})',
            r'(\d{4}-\d{2}-\d{2}).*?to.*?(\d{4}-\d{2}-\d{2})'
        ]
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.lower()
            for pattern in date_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    if 'start' in pattern:
                        start_date = match.group(1)
                    elif 'end' in pattern:
                        end_date = match.group(1)
                    elif len(match.groups()) == 2:  # Duration pattern
                        start_date, end_date = match.groups()
        
        return start_date, end_date
    
    @staticmethod
    def extract_media_plan_table(doc: Document) -> List[Dict[str, Any]]:
        """Extract media plan table data from Word document"""
        media_plan_data = []
        
        for table in doc.tables:
            # Look for tables that might contain media plan data
            if len(table.rows) < 2:
                continue
                
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            
            # Check if this looks like a media plan table
            if any(keyword in ' '.join(headers) for keyword in ['platform', 'channel', 'cost', 'budget', 'objective']):
                for row in table.rows[1:]:  # Skip header row
                    row_data = {}
                    for i, cell in enumerate(row.cells):
                        if i < len(headers):
                            header = headers[i]
                            value = cell.text.strip()
                            
                            # Parse numeric values for cost/budget
                            if 'cost' in header or 'budget' in header or 'spend' in header:
                                # Extract numeric value from currency strings
                                numeric_match = re.search(r'[\d,]+\.?\d*', value.replace(',', ''))
                                if numeric_match:
                                    try:
                                        row_data[header] = float(numeric_match.group().replace(',', ''))
                                    except ValueError:
                                        row_data[header] = value
                                else:
                                    row_data[header] = value
                            else:
                                row_data[header] = value
                    
                    if row_data and any(row_data.values()):  # Skip empty rows
                        media_plan_data.append(row_data)
        
        return media_plan_data
    
    @staticmethod
    def extract_creative_checklist(doc: Document) -> List[str]:
        """Extract creative requirements checklist from Word document"""
        checklist_items = []
        in_checklist = False
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Look for creative checklist section
            if 'creative' in text.lower() and ('checklist' in text.lower() or 'requirements' in text.lower()):
                in_checklist = True
                continue
            
            # Stop at next major section
            if in_checklist and text and text[0].isupper() and ':' in text and 'creative' not in text.lower():
                break
            
            # Extract checklist items (lines with checkmarks or bullet points)
            if in_checklist and (text.startswith('✓') or text.startswith('✔') or text.startswith('•') or text.startswith('-')):
                item = re.sub(r'^[✓✔•\-\s]+', '', text).strip()
                if item:
                    checklist_items.append(item)
        
        return checklist_items
    
    @staticmethod
    def extract_strategy_channels(doc: Document) -> List[str]:
        """Extract channel names from strategy explainer section after 'Channel:' labels"""
        strategy_channels = []
        
        # Process all paragraphs and table cells to find "Channel:" patterns
        all_text_elements = []
        
        # Add paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                all_text_elements.append(para.text.strip())
        
        # Add table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text_elements.append(cell.text.strip())
        
        # Look for "Channel:" pattern in all text elements
        channel_pattern = r'channel:\s*([^\n\r]+)'
        
        for text in all_text_elements:
            # Search for channel patterns (case insensitive)
            matches = re.finditer(channel_pattern, text, re.IGNORECASE)
            for match in matches:
                channel_name = match.group(1).strip()
                
                # Clean up the channel name (remove extra punctuation, etc.)
                channel_name = re.sub(r'^[•\-\*\s]+', '', channel_name)  # Remove bullet points
                channel_name = re.sub(r'[.,:;]+$', '', channel_name)     # Remove trailing punctuation
                channel_name = channel_name.strip()
                
                if channel_name and channel_name not in strategy_channels:
                    strategy_channels.append(channel_name)
        
        return strategy_channels


class CurrencyExtractor:
    """Utility class for extracting currency amounts from documents"""
    
    @staticmethod
    def extract_currency_amounts_from_text(doc: Document) -> List[Dict[str, Any]]:
        """Extract currency amounts from document text, but stop at strategy explainer section"""
        currency_data = []
        
        # Process all paragraphs and table cells
        all_text_elements = []
        
        # Add paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                all_text_elements.append(para.text.strip())
        
        # Add table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text_elements.append(cell.text.strip())
        
        # Look for currency amounts in specific sections:
        # 1. Before "2. Strategy Explainer" 
        # 2. Skip the strategy explainer narrative text
        # 3. Include amounts from tables that appear after strategy (these are plan data)
        
        strategy_explainer_index = None
        creative_checklist_index = None
        
        # First pass: find section boundaries
        for i, text in enumerate(all_text_elements):
            text_lower = text.lower()
            
            # Find strategy explainer section
            for pattern in STRATEGY_EXPLAINER_PATTERNS:
                if re.search(pattern, text_lower):
                    strategy_explainer_index = i
                    break
            
            # Find creative checklist section (marks end of strategy narrative)
            if 'creative' in text_lower and ('checklist' in text_lower or 'requirements' in text_lower):
                creative_checklist_index = i
                break
        
        # Second pass: extract currency amounts from appropriate sections
        for i, text in enumerate(all_text_elements):
            text_lower = text.lower()
            
            # Determine if we should process this element for currency amounts
            should_process = False
            
            if strategy_explainer_index is None:
                # No strategy explainer found, process all elements
                should_process = True
            elif i < strategy_explainer_index:
                # Before strategy explainer - process
                should_process = True  
            elif creative_checklist_index is not None and i >= creative_checklist_index:
                # After creative checklist (likely contains plan tables) - process
                should_process = True
            # Skip elements between strategy explainer and creative checklist (strategy narrative)
            
            if not should_process:
                continue
            
            # Find all currency amounts in this text
            for pattern in CURRENCY_PATTERNS:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.group()
                    
                    # Extract numeric value
                    numeric_match = re.search(r'[\d,]+\.?\d*', amount_str.replace(',', ''))
                    if numeric_match:
                        try:
                            amount = float(numeric_match.group().replace(',', ''))
                            
                            # Look for context in current and surrounding text
                            context_text = text_lower
                            
                            # Check previous and next elements for context
                            if i > 0:
                                context_text = all_text_elements[i-1].lower() + " " + context_text
                            if i < len(all_text_elements) - 1:
                                context_text = context_text + " " + all_text_elements[i+1].lower()
                            
                            # Identify platform/channel
                            platform = None
                            for keyword in PLATFORM_KEYWORDS:
                                if keyword in context_text:
                                    platform = keyword
                                    break
                            
                            # Check if near objective or impression data
                            near_objective = any(obj in context_text for obj in OBJECTIVE_KEYWORDS)
                            near_impressions = 'impression' in context_text or 'reach' in context_text
                            
                            currency_data.append({
                                'amount': amount,
                                'original_text': amount_str,
                                'context': text[:100] + "..." if len(text) > 100 else text,
                                'platform': platform,
                                'near_objective': near_objective,
                                'near_impressions': near_impressions,
                                'before_strategy': True  # All amounts are now before strategy section
                            })
                            
                        except ValueError:
                            continue
        
        return currency_data


def normalize_channel_name(channel: str) -> str:
    """Normalize channel names for comparison"""
    channel = channel.lower().strip()
    return CHANNEL_NORMALIZATIONS.get(channel, channel)